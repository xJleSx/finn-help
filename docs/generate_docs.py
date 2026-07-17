#!/usr/bin/env python3
"""Generate FinAdvisor documentation (.docx).

Run: uv run python docs/generate_docs.py
Requires: python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime
import subprocess


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elem)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
        set_cell_shading(cell, 'D9E2F3')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.5)
    pf.right_indent = Cm(0.5)
    pPr = p._element.get_or_add_pPr()
    shd = pPr.makeelement(qn('w:shd'), {
        qn('w:fill'): 'F0F0F0',
        qn('w:val'): 'clear',
    })
    pPr.append(shd)
    for i, line in enumerate(text.split('\n')):
        if i > 0:
            bp = doc.add_paragraph()
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(0)
            pf2 = bp.paragraph_format
            pf2.left_indent = Cm(0.5)
            pf2.right_indent = Cm(0.5)
            pPr2 = bp._element.get_or_add_pPr()
            shd2 = pPr2.makeelement(qn('w:shd'), {
                qn('w:fill'): 'F0F0F0',
                qn('w:val'): 'clear',
            })
            pPr2.append(shd2)
            run = bp.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            continue
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(8)
    doc.add_paragraph('')


def build():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    # ── Title ──
    for _ in range(4):
        doc.add_paragraph('')
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('FinAdvisor\nAI-финансовый ассистент для MOEX')
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    doc.add_paragraph('')
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Техническая документация')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

    doc.add_paragraph('')
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = dp.add_run(datetime.date.today().strftime('%d.%m.%Y'))
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── TOC ──
    doc.add_heading('Содержание', level=1)
    toc = [
        '1. Архитектура',
        '2. Сбор данных (коллекторы)',
        '3. Технический анализ',
        '   3.1. RSI',
        '   3.2. MACD',
        '   3.3. Bollinger Bands',
        '   3.4. SMA, ATR',
        '   3.5. Сигнальная система',
        '4. Фундаментальный анализ',
        '   4.1. Акции: секторная нормализация',
        '   4.2. Multiplicative risk',
        '   4.3. Облигации: YTM, спред, рейтинг',
        '5. ML-модуль',
        '   5.1. Признаки',
        '   5.2. Целевая переменная',
        '   5.3. XGBoost / LightGBM / CatBoost',
        '   5.4. Ансамбль + Stacking',
        '   5.5. Веса моделей по OOS',
        '   5.6. Walk-forward',
        '   5.7. Trend Predictor',
        '   5.8. Неопределённость',
        '   5.9. Price Targets',
        '6. Волатильность (VolatilityRegimeDetector)',
        '7. Мультивременной анализ (MTF)',
        '8. Сентимент',
        '9. Геополитический риск',
        '10. Fusion Engine',
        '   10.1. Веса',
        '   10.2. Risk-профили',
        '   10.3. Динамические коррекции',
        '   10.4. Макро',
        '   10.5. Bond-специфичный scoring',
        '   10.6. Итоговый сигнал',
        '11. Алерты',
        '12. Торговля',
        '13. Портфель',
        '14. Уведомления',
        '15. LLM',
        '16. CLI и API',
        '17. База данных',
        '18. Социальный сентимент',
        '19. Детекция аномалий',
        '20. Планировщик',
        '21. Web-фронтенд',
        '22. Контракты данных',
        '23. Деградация',
        '24. Известные проблемы',
        '25. Пример сквозной',
        '26. Таблицы порогов',
    ]
    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)

    doc.add_page_break()

    # ═══════════════ 1. АРХИТЕКТУРА ═══════════════
    doc.add_heading('1. Архитектура', level=1)

    doc.add_paragraph(
        'FastAPI (Python 3.13) + Next.js 16 (TypeScript). '
        '24 пакета, 93 Python-модуля, 83 тестовых файла. '
        'Рынок: MOEX (акции, облигации).'
    )

    doc.add_paragraph('Pipeline анализа (каждый шаг опциональный, Fusion переживёт отсутствие любого модуля):')
    p = doc.add_paragraph()
    r = p.add_run('1. DataLoader -> ')
    r = p.add_run('2. TechnicalAnalyzer -> ')
    r = p.add_run('3. FundamentalAnalyzer -> ')
    r = p.add_run('4. ML Ensemble (XGB+LGB+Cat) -> ')
    r = p.add_run('5. MultiTimeframe -> ')
    r = p.add_run('6. VolatilityRegime -> ')
    r = p.add_run('7. GeoRisk -> ')
    r = p.add_run('8. Sentiment -> ')
    r = p.add_run('9. EventContext -> ')
    r = p.add_run('10. Fusion Engine -> ')
    r = p.add_run('11. LLM Advice (опционально)')
    r.font.size = Pt(9)

    doc.add_paragraph('Модули src/:')
    doc.add_paragraph(
        'collectors/    — данные: MOEX, ЦБ, новости, соцсети\n'
        'analysis/      — технический, фундаментальный, ML, MTF, сценарии, аномалии\n'
        'signals/       — Fusion Engine\n'
        'alerts/        — алерты с дедупликацией\n'
        'trading/       — брокеры, ордера, риск, комплаенс, налоги\n'
        'portfolio/     — аллокатор, риск-метрики\n'
        'notifications/ — email, telegram, webpush\n'
        'llm/           — Groq + Ollama, промпты\n'
        'interfaces/    — FastAPI, Telegram, NLQ\n'
        'core/          — DI, auth (JWT/TOTP), structlog, tracing, resilience\n'
        'db/            — SQLAlchemy async, 29 миграций\n'
        'scheduler/     — Celery\n'
        'social/        — сентимент соцсетей\n'
        'geo/           — геополитический риск\n'
        'data/          — пайплайн новостей и событий\n'
        'cli/           — Typer команды\n'
        'reports/       — PDF\n'
    )

    doc.add_heading('AnalysisService (src/analysis/service.py)', level=2)
    doc.add_paragraph(
        'Оркестратор, вызывает все анализаторы по очереди. '
        'Не требует всех модулей — каждый возвращает None при недоступности, '
        'Fusion Engine его переживёт.'
    )
    code(doc, '''class AnalysisService:
    def __init__(self):
        self.analyzer = TechnicalAnalyzer()
        self.fundamental = FundamentalAnalyzer()
        self.fusion = SignalFusionEngine()
        self.volatility = VolatilityRegimeDetector()
        self.mtf = MultiTimeframeAnalyzer()
        self.ml = MLCoordinator()
        self.events = EventFeatureBuilder()

    async def analyze(self, ticker: str, db, user_id=None) -> dict:
        prices, inds, divs, metrics, events = await self.loader.load(ticker, db)
        df = self._price_df(prices)
        tech = self.analyzer.compute_all(df)
        tech_sig = self.analyzer.generate_signal(tech) if not df.empty else None
        fund = self.fundamental.analyze(df, self._dividend_df(divs), metrics)
        ml_result = await self._compute_ml(df, ind_df, ticker, events)
        vol = self.volatility.detect(df)
        mtf_result = self.mtf.analyze(df)
        geo = await self._compute_geo(ticker, db)
        sent = await self._compute_sentiment(ticker)
        macro = await self.loader.load_macro(db)
        event_ctx = self.events.build(events)

        fused = self.fusion.fuse(
            ticker=ticker, technical=tech_sig, fundamental=fund,
            geo=geo, ml_prediction=ml_result, volatility_regime=vol,
            macro_context=macro, sentiment=sent, mtf=mtf_result,
            event_context=event_ctx, user_id=user_id,
        )
        if user_id:
            fused["llm_advice"] = await llm.generate_advice(fused)
        return fused''')

    doc.add_page_break()

    # ═══════════════ 2. СБОР ДАННЫХ ═══════════════
    doc.add_heading('2. Сбор данных (коллекторы)', level=1)
    doc.add_paragraph('src/collectors/')

    doc.add_paragraph(
        'BaseCollector задаёт интерфейс fetch() с кэшем (TTLCache + Redis) '
        'и retry через tenacity (exponential backoff: 1s-60s, 5 попыток). '
        'CircuitBreaker: при 5+ ошибках подряд отключает источник на 30 секунд.'
    )
    code(doc, '''class BaseCollector(ABC):
    def __init__(self, cache_ttl=300):
        self.cache = TTLCache(maxsize=100, ttl=cache_ttl)

    @abstractmethod
    async def fetch(self, **kwargs):
        ...

    async def safe_fetch(self, **kwargs):
        if not self.breaker.ready():
            raise CircuitBreakerOpenError
        try:
            result = await retry_call(self.fetch, tries=5, backoff=2)
            self.breaker.succeed()
            return result
        except Exception:
            self.breaker.fail()
            raise''')

    add_table(doc,
        ['Файл', 'Источник', 'Что отдаёт'],
        [
            ('moex.py', 'MOEX ISS', 'OHLCV, индикаторы, дивиденды, стакан'),
            ('bonds.py', 'MOEX', 'облигации, купоны, даты, номинал'),
            ('cbr.py', 'ЦБ РФ', 'ключевая ставка, курсы, CPI'),
            ('macro.py', 'агрегация', 'Brent, M2, IMOEX, ОФЗ'),
            ('news.py', 'RSS', 'TASS, RBC, Interfax — до 5 статей/фид'),
            ('social.py', 'Telegram', 'посты каналов'),
            ('financials.py', 'MOEX', 'отчётность МСФО/РСБУ'),
            ('fundamental.py', 'MOEX', 'P/E, P/B, ROE, D/E'),
            ('profiles.py', 'MOEX', 'описание, сектор'),
        ]
    )

    doc.add_paragraph(
        'Подводные камни: MOEX ISS отдаёт данные порциями по 100 записей — '
        'нужна пагинация (есть в moex.py). ЦБ РФ блокирует при 10+ запросах/сек — '
        'встроен rate limit. NewsCollector падает при недоступности RSS-ленты — '
        'обрабатывается как пустой результат.'
    )

    doc.add_page_break()

    # ═══════════════ 3. ТЕХНИЧЕСКИЙ АНАЛИЗ ═══════════════
    doc.add_heading('3. Технический анализ', level=1)
    doc.add_paragraph('src/analysis/technical.py | TechnicalAnalyzer')

    doc.add_paragraph(
        'Вход: DataFrame с колонками date, open, high, low, close, volume. '
        'compute_all() добавляет колонки-индикаторы по очереди.'
    )
    code(doc, '''class TechnicalAnalyzer:
    def compute_all(self, df):
        if df.empty: return df
        df = df.sort_values("date").copy()
        for period in [20, 50, 200]:
            df = self.sma(df, period)
        df = self.rsi(df, 14)
        df = self.macd(df)
        df = self.bollinger_bands(df, 20)
        df = self.volume_sma(df, 20)
        return self.atr(df, 14)''')

    doc.add_heading('3.1. RSI', level=2)
    doc.add_paragraph('Wilder RSI(14). Отличие от стандартного pandas: alpha=1/14, а не ewm(span=14).')
    code(doc, '''def rsi(self, df, period=14):
    delta = df["close"].diff()
    gain = delta.where(delta > 0, 0.0)
    loss = (-delta).where(delta < 0, 0.0)
    alpha = 1.0 / period  # Wilder, не 2/(period+1)
    avg_gain = gain.ewm(alpha=alpha, adjust=False).mean()
    avg_loss = loss.ewm(alpha=alpha, adjust=False).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    df["rsi"] = 100 - (100 / (1 + rs))
    df["rsi"] = np.where(avg_loss == 0,
        np.where(avg_gain == 0, 50.0, 100.0), df["rsi"])
    return df''')

    doc.add_paragraph(
        'Почему не ewm(span=14): Wilder RSI использует α=1/N, а EMA с span=N '
        'даёт α=2/(N+1) — разные веса. Замена была багом — RSI сглаживался сильнее '
        'и давал меньше пересечений 30/70.\n'
        'avg_loss==0 AND avg_gain==0 → RSI=50 (флэт). '
        'avg_loss==0 AND avg_gain>0 → RSI=100 (только рост). '
        'replace(0, NaN) защищает от деления на ноль.'
    )

    doc.add_heading('3.2. MACD', level=2)
    code(doc, '''def macd(self, df):
    ema_12 = df["close"].ewm(span=12, adjust=False).mean()
    ema_26 = df["close"].ewm(span=26, adjust=False).mean()
    df["macd_line"] = ema_12 - ema_26
    df["macd_signal"] = df["macd_line"].ewm(span=9, adjust=False).mean()
    df["macd_hist"] = df["macd_line"] - df["macd_signal"]
    return df''')
    doc.add_paragraph(
        'Стандартный MACD(12,26,9). В scoring смотрится пересечение '
        'гистограммой нуля с порогом max(0.01, std(macd_hist)*0.1) — '
        'фильтр whipsaw. Порог эмпирический: на VTBR std высокий, порог '
        'может быть слишком большим — сигналы пропускаются.'
    )

    doc.add_heading('3.3. Bollinger Bands', level=2)
    code(doc, '''def bollinger_bands(self, df, period=20):
    df["bb_mid"] = df["close"].rolling(window=period).mean()
    std = df["close"].rolling(window=period).std()
    df["bb_upper"] = df["bb_mid"] + std * 2
    df["bb_lower"] = df["bb_mid"] - std * 2
    return df''')
    doc.add_paragraph('20 периодов, 2σ. Никакой модификации — канон Боллинджера.')

    doc.add_heading('3.4. SMA и ATR', level=2)
    code(doc, '''def sma(self, df, period):
    df[f"sma_{period}"] = df["close"].rolling(window=period).mean()
    return df

def atr(self, df, period=14):
    hl = df["high"] - df["low"]
    hc = (df["high"] - df["close"].shift()).abs()
    lc = (df["low"] - df["close"].shift()).abs()
    tr = pd.concat([hl, hc, lc], axis=1).max(axis=1)
    df["atr"] = tr.ewm(alpha=1.0/period, adjust=False).mean()
    return df''')
    doc.add_paragraph(
        'SMA: 20 (месяц), 50 (квартал), 200 (год). '
        'ATR — Wilder ATR, не SMA: ewm(alpha=1/14), а не rolling(14).mean(). '
        'Разница: SMA даёт равные веса, EWM — больший вес свежим данными.'
    )

    doc.add_heading('3.5. Сигнальная система', level=2)
    doc.add_paragraph(
        'generate_signal() собирает score из всех индикаторов и нормирует на max_score. '
        "Порог 0.20 — обоснован walk-forward'ом на 10 MOEX-тикерах + 7 синтетик."
    )
    code(doc, '''def generate_signal(self, df):
    if df.empty or len(df) < 50:
        return {"action": "NEUTRAL", "confidence": 0.0,
                "score": 0.0, "reasons": ["недостаточно данных"]}

    latest = df.iloc[-1]
    score, max_score = 0.0, 0.0
    reasons = []

    # RSI: перекупленность/перепроданность
    if not pd.isna(latest.get("rsi")):
        max_score += 1.0
        if latest["rsi"] < 30:      score += 1.0; reasons.append("RSI перепродан")
        elif latest["rsi"] > 70:    score -= 1.0; reasons.append("RSI перекуплен")

    # MACD гистограмма пересекает ноль (с whipsaw-защитой)
    if not pd.isna(latest.get("macd_hist")):
        max_score += 1.0
        prev = df.iloc[-2]
        threshold = max(0.01, df["macd_hist"].std() * 0.1)
        if latest["macd_hist"] > threshold and prev["macd_hist"] <= 0:
            score += 1.0; reasons.append("MACD бычий")
        elif latest["macd_hist"] < -threshold and prev["macd_hist"] >= 0:
            score -= 1.0; reasons.append("MACD медвежий")

    # SMA: цена над/под скользящими (по 0.5 балла за каждую)
    for col in ["sma_20", "sma_50", "sma_200"]:
        val = latest.get(col)
        if pd.isna(val):  continue  # мало данных — не штрафуем
        max_score += 0.5
        if latest["close"] > val:    score += 0.5
        elif latest["close"] < val:  score -= 0.5

    # Bollinger: касание границ
    if all(not pd.isna(latest.get(k)) for k in ("close","bb_lower","bb_upper")):
        max_score += 0.5
        if latest["close"] <= latest["bb_lower"]:    score += 0.5
        elif latest["close"] >= latest["bb_upper"]:  score -= 0.5

    # Импульс: объём, ATR, однодневная доходность
    score, max_score = self._add_momentum_scores(df, score, max_score, reasons)

    normalized = score / max_score if max_score > 0 else 0.0
    if normalized > 0.20:    action = "BUY"
    elif normalized < -0.20:  action = "SELL"
    else:                     action = "HOLD"

    return {"action": action, "confidence": abs(normalized),
            "score": normalized, "reasons": reasons}''')

    doc.add_paragraph('_add_momentum_scores() — три дополнительных компонента:')
    code(doc, '''def _add_momentum_scores(self, df, score, max_score, reasons):
    latest = df.iloc[-1]
    vol_ratio = latest.get("volume_ratio")  # объём / SMA20 объёма
    if not pd.isna(vol_ratio):
        max_score += 0.3
        if vol_ratio > 1.5:
            score += 0.3
            reasons.append("Аномальный объём")

    atr = latest.get("atr_pct")
    if not pd.isna(atr):
        max_score += 0.2
        if atr < df["atr_pct"].mean() * 0.7:
            reasons.append("Низкая волатильность — возможен выход")

    if len(df) >= 2:
        ret_1d = (latest["close"] / df.iloc[-2]["close"]) - 1
        max_score += 0.3
        if ret_1d > 0.01:    score += 0.15
        elif ret_1d < -0.01:  score -= 0.15
    return score, max_score''')
    doc.add_paragraph(
        'volume_ratio = объём / SMA20_объёма. Порог 1.5 — эмпирический. '
        'atr_pct = ATR / close. ret_1d > 1% даёт 0.15 балла.'
    )

    doc.add_page_break()

    # ═══════════════ 4. ФУНДАМЕНТАЛЬНЫЙ ═══════════════
    doc.add_heading('4. Фундаментальный анализ', level=1)
    doc.add_paragraph('src/analysis/fundamental.py | FundamentalAnalyzer')

    doc.add_heading('4.1. Акции: секторная нормализация', level=2)
    doc.add_paragraph(
        'Вместо абсолютных порогов (P/E>30, P/B>5) используем секторные медианы. '
        'Параметр sector передаётся в analyze(), по умолчанию "Прочее". '
        'Если сектор неизвестен — используется fallback.'
    )
    code(doc, '''SECTOR_BENCHMARKS = {
    "Финансы":         {"pe_median": 6.0,  "pb_median": 0.8,  "roe_median": 15.0, "de_median": 5.0},
    "Нефть":           {"pe_median": 5.0,  "pb_median": 1.0,  "roe_median": 18.0, "de_median": 1.5},
    "Металлы":         {"pe_median": 7.0,  "pb_median": 1.5,  "roe_median": 20.0, "de_median": 1.0},
    "IT":              {"pe_median": 25.0, "pb_median": 4.0,  "roe_median": 10.0, "de_median": 0.8},
    "Потребительский": {"pe_median": 10.0, "pb_median": 2.0,  "roe_median": 15.0, "de_median": 1.5},
    "Телеком":         {"pe_median": 8.0,  "pb_median": 1.2,  "roe_median": 12.0, "de_median": 2.5},
    "Энергетика":      {"pe_median": 6.0,  "pb_median": 1.0,  "roe_median": 10.0, "de_median": 3.0},
    "Химия":           {"pe_median": 9.0,  "pb_median": 1.8,  "roe_median": 18.0, "de_median": 1.2},
    "Транспорт":       {"pe_median": 8.0,  "pb_median": 1.5,  "roe_median": 12.0, "de_median": 2.0},
    "Прочее":          {"pe_median": 10.0, "pb_median": 1.5,  "roe_median": 12.0, "de_median": 1.5},
}
PE_HIGH_MULTIPLE = 3.0   # P/E > median*3 → anomaly
ROE_LOW_MULTIPLE = 0.3   # ROE < median*0.3 → anomaly
DE_HIGH_MULTIPLE = 2.0   # D/E > median*2 → anomaly''')
    doc.add_paragraph(
        'Например, для IT-сектора: P/E > 25*3 = 75 → аномалия. '
        'Для Нефти: P/E > 5*3 = 15 → аномалия. '
        'Значения медиан — 2024-2025 данные MOEX. Нуждаются в '
        'периодическом пересмотре (см. раздел 24).'
    )

    doc.add_heading('4.2. Multiplicative risk', level=2)
    doc.add_paragraph(
        'Вместо min(risk_score, 1.0) — мультипликативная модель. '
        'Зачем: при 3 аномалиях по 0.25 аддитивный метод даёт 0.75, '
        'при 6 — всё ещё 1.0 (нет дифференциации). '
        'Мультипликативный: 3→0.578, 6→0.822.'
    )
    code(doc, '''def _multiplicative_risk(risks):
    total = 1.0
    for r in risks:
        total *= max(0.0, 1.0 - r)
    return 1.0 - total''')
    doc.add_paragraph('Формула: 1 - ∏(1 - rᵢ). Риски предполагаются независимыми.')

    doc.add_heading('4.3. Облигации: YTM, спред, рейтинг', level=2)
    doc.add_paragraph(
        'analyze_bond() оценивает риск по спреду к ключевой ставке, '
        'дюрации и кредитному рейтингу. Две шкалы рейтинга: '
        'международная (S&P) и российская (АКРА/Эксперт РА).'
    )
    code(doc, '''def analyze_bond(self, bond_offering, key_rate=None):
    ytm = bond_offering.get("yield_to_maturity")
    credit_rating = bond_offering.get("credit_rating")
    duration = bond_offering.get("duration_years")

    # Спред — основной метод
    if ytm is not None and key_rate is not None:
        spread = ytm - key_rate
        if spread > 5:      risk_parts.append(0.25)  # дефолтный риск
        elif spread > 3:    risk_parts.append(0.10)  # повышенный
        elif spread < -2:   risk_parts.append(0.05)  # аномалия

    # Fallback: абсолютный YTM (нет ключевой ставки)
    if ytm is not None and (key_rate is None or key_rate <= 0):
        if ytm > 25:    risk_parts.append(0.30)
        elif ytm > 15:  risk_parts.append(0.10)

    if duration is not None:
        if duration > 10:     risk_parts.append(0.20)
        elif duration > 5:    risk_parts.append(0.10)

    rating_risk = _rating_risk(credit_rating)
    risk_parts.append(rating_risk)
    return {"risk": _multiplicative_risk(risk_parts),
            "anomalies": anomalies, "signals": signals}''')

    doc.add_paragraph('Маппинг рейтингов (префиксный матчинг):')
    code(doc, '''def _rating_risk(rating):
    if not rating: return 0.1
    r = rating.strip().lower().replace(" ", "")
    for risk, prefixes in NATIONAL_MAP + INTERNATIONAL_MAP:
        if any(r.startswith(p) for p in prefixes):
            return risk
    return 0.1  # неизвестный рейтинг — mild penalty

# NATIONAL_MAP: ruAAA→0, ruAA+→0, ..., ruD→0.30
# INTERNATIONAL_MAP: AAA→0, AA+→0, ..., D→0.30''')
    doc.add_paragraph(
        'Подводный камень: "BB+" — международный рейтинг (0.20), '
        '"ruBB+" — российский (0.20 совпало, но не всегда). '
        'Префикс "ru" отличает национальную шкалу от международной. '
        'Если рейтинг не распознан — 0.1 (не 0 — консервативно).'
    )

    doc.add_page_break()

    # ═══════════════ 5. ML ═══════════════
    doc.add_heading('5. ML-модуль', level=1)
    doc.add_paragraph('src/analysis/ml/')

    doc.add_heading('5.1. Признаки', level=2)
    doc.add_paragraph(
        'prepare_features() из OHLCV + индикаторов делает 22+ признака '
        '(14 базовых + 4 событийных + 7 макро).'
    )
    code(doc, '''BASE_FEATURES = ["close", "rsi", "macd_hist", "sma_20", "sma_50",
    "price_sma20", "price_sma50", "sma20_sma50",
    "rsi_norm", "macd_signal_binary", "atr_pct",
    "volume_ratio", "bb_width", "hist_vol_20"]
EVENT_FEATURES = ["event_count_30d", "event_severity_30d",
                   "sanctions_30d", "days_since_major_event"]
MACRO_FEATURES = ["brent", "key_rate", "usd_rate", "imoex", "cpi", "ofz_10y"]

def prepare_features(df):
    result = df[["rsi","macd_hist","sma_20","sma_50","close"]].copy()
    result["price_sma20"] = result["close"] / result["sma_20"]
    result["price_sma50"] = result["close"] / result["sma_50"]
    result["sma20_sma50"] = result["sma_20"] / result["sma_50"]
    result["rsi_norm"] = result["rsi"] / 100
    result["macd_signal_binary"] = (result["macd_hist"] > 0).astype(int)
    result["atr_pct"] = atr / close
    result["volume_ratio"] = volume / volume_sma_20
    result["bb_width"] = (bb_up - bb_low) / bb_mid
    result["hist_vol_20"] = returns.rolling(20).std()
    for c in EVENT_FEATURES + MACRO_FEATURES:
        result[c] = df[c] if c in df.columns else 0
    return result.dropna()''')
    doc.add_paragraph(
        'price_sma20 = close/SMA20 (>1 = выше средней). '
        'macd_signal_binary = 1 если MACD_hist > 0. '
        'Если макро/событий нет в df — заполняем нулём (ML переживёт).'
    )

    doc.add_heading('5.2. Целевая переменная', level=2)
    doc.add_paragraph('Бинарная: 1 = BUY, 0 = SELL, NaN = шум (не учимся).')
    code(doc, '''def build_labels(close_series, lookahead=5, threshold=0.03):
    future_ret = close_series.shift(-lookahead) / close_series - 1
    y = np.where(future_ret > threshold, 1,
         np.where(future_ret < -threshold, 0, np.nan))
    return y, ~np.isnan(y)

def compute_threshold(close_series, lookahead=5, fallback=0.03):
    ret = close_series.pct_change().dropna()
    if len(ret) < lookahead + 1: return fallback
    return max(fallback, ret.std() * 0.5)''')
    doc.add_paragraph(
        'Порог 3% — примерно 2× среднедневной спред MOEX. '
        'Адаптивный: при высокой волатильности порог растёт, '
        'меньше шумовых меток. lookahead=5 дней — краткосрок, '
        'для долгосрочных стратегий нужно менять.',
    )

    doc.add_heading('5.3. Модели', level=2)
    doc.add_paragraph(
        'Три модели наследуют BaseMLClassifier. Разница только в _create_model().'
    )
    code(doc, '''class XGBoostClassifier(BaseMLClassifier):
    def _create_model(self):
        return xgb.XGBClassifier(eval_metric="logloss", verbosity=0)

class LightGBMClassifier(BaseMLClassifier):
    def _create_model(self):
        return lgb.LGBMClassifier(verbosity=-1, deterministic=True)

class CatBoostClassifierModel(BaseMLClassifier):
    def _create_model(self):
        return CatBoostClassifier(
            iterations=n_estimators, max_depth=max_depth,
            learning_rate=learning_rate, verbose=0,
            allow_writing_files=False)''')
    doc.add_paragraph(
        'Гиперпараметры (n_estimators, max_depth, learning_rate) — '
        'из config.py (Settings.ml_*). Общие для всех трёх, но CatBoost '
        'иногда требует отдельной настройки — следите за переобучением.'
    )

    doc.add_heading('5.4. Ансамбль + Stacking', level=2)
    doc.add_paragraph(
        'EnsemblePredictor.predict(): '
        '(1) все три модели предсказывают, (2) веса по OOS accuracy, '
        '(3) взвешенное среднее probability, (4) majority vote, '
        '(5) meta-learner (LogisticRegression на OOF) — если accuracy ≥ 52%.'
    )
    code(doc, '''def predict(self, df, anomaly_mask=None):
    models = [("xgb", self.xgb), ("lgb", self.lgb), ("cat", self.cat)]
    for name, model in models:
        pred = model.predict(df, anomaly_mask=anomaly_mask)
        oos = self._walk_forward_validate(df, model)
        named_oos[name] = oos
    weights = self._get_weights([named_oos[n] for n in ("xgb","lgb","cat")])
    avg_prob = sum(r["probability"] * w for r, w in zip(results, weights)) / total_w
    meta_probs = self._stacking_predict(df, results)
    if meta_probs is not None:
        avg_prob = meta_probs
    uncertainty = np.std([r["probability"] for r in results]) * 2
    return {"action": action, "confidence": final_confidence,
            "probability": avg_prob, "uncertainty": uncertainty}''')

    doc.add_paragraph('Meta-learner:')
    code(doc, '''def _train_meta_oof(self, df):
    features = prepare_features(df)
    y, mask = build_labels(df["close"])
    x_all, y_all = features[mask], y[mask].astype(int)
    splits = temporal_split(len(x_all))
    x_train, y_train = x_all[splits["train"]], y_all[splits["train"]]
    x_val, y_val     = x_all[splits["val"]],   y_all[splits["val"]]
    oof_probs = []
    for name in ("xgb", "lgb", "cat"):
        m = getattr(self, name)
        m.fit(x_train, y_train)
        oof_probs.append(m._model.predict_proba(x_val)[:, 1])
    meta_x = StandardScaler().fit_transform(np.column_stack(oof_probs))
    meta = LogisticRegression(max_iter=2000, C=0.5, random_state=42)
    meta.fit(meta_x, y_val)
    if meta.score(meta_x, y_val) < 0.52:
        self._meta_model = None  # отбрасываем, meta хуже random''')
    doc.add_paragraph(
        'C=0.5 — L2-регуляризация. Discard 52% — эмпирика: meta должна '
        'хоть немного превосходить random. TTL кэша meta — 24 часа.'
    )

    doc.add_heading('5.5. Веса моделей по OOS', level=2)
    code(doc, '''def model_weight_from_oos(oos):
    acc = oos.get("oos_accuracy", 0.5)
    folds = oos.get("folds_completed", 0)
    if folds == 0 or acc < 0.52: return 0.0
    return (acc - 0.5) * 4 * min(folds / 3, 1)''')
    doc.add_paragraph('0.5 → 0.0, 0.6 → 0.4, 0.75 → 1.0. Линейно.')

    doc.add_heading('5.6. Walk-forward', level=2)
    code(doc, '''def walk_forward_validate(model, x, y, n_splits=3, min_train=60):
    n, fold_size = len(x), (n - min_train) // n_splits
    accs = []
    for i in range(n_splits):
        test_end = n - i * fold_size
        test_start = test_end - fold_size
        train_end = test_start - 20  # gap!
        model.fit(x[:train_end], y[:train_end])
        preds = model.predict(x[test_start:test_end])
        accs.append(np.mean(preds == y[test_start:test_end]))
    return {"oos_accuracy": np.mean(accs), "folds_completed": len(accs)}''')
    doc.add_paragraph(
        'Gap=20 дней между train/test — защита от автокорреляции цен. '
        'Без gap модель «подглядывала» бы в будущее через лаговые признаки. '
        '3-fold: достаточно для временных рядов (Pesaran-Timmermann).'
    )

    doc.add_heading('5.7. Trend Predictor', level=2)
    doc.add_paragraph(
        'StatsModelsTrendPredictor — замена Prophet. ExponentialSmoothing '
        '(сезонность 365d при 2+ годах данных, иначе 7d). Fallback — линейная регрессия.'
    )
    code(doc, '''def _fit(self, df):
    trend = df[["date","close"]].copy()
    trend.columns = ["ds", "y"]
    if len(trend) >= 730:
        model = ExponentialSmoothing(y, trend="add",
                     seasonal="add", seasonal_periods=365)
    elif len(trend) >= 100:
        model = ExponentialSmoothing(y, trend="add",
                     seasonal="add", seasonal_periods=7)
    else:
        return self._linear_fallback(trend)
    fitted = model.fit()
    forecast = fitted.forecast(days_ahead)
    residual_std = np.std(y - fitted.fittedvalues)
    return {"target_price": forecast[-1], "trend_slope": tanh(slope/mean*100),
            "trend_strength": r2_score(actual, predicted)}''')
    doc.add_paragraph(
        'Prophet убрали из-за проблем с зависимостями (pystan). '
        'ExponentialSmoothing даёт сопоставимое качество на MOEX. '
        'Доверительный интервал: forecast ± 1.96×σ_residual.'
    )

    doc.add_heading('5.8. Неопределённость', level=2)
    doc.add_paragraph('Два уровня:')
    code(doc, '''# 1. Bootstrap базовых моделей
def _bootstrap_interval(self, features, proba):
    preds = [bm.predict_proba(features.iloc[-1:])[0,1]
             for bm in self._bootstrap_models]
    return np.percentile(preds, 5), np.percentile(preds, 95)

# 2. Разброс в ансамбле
uncertainty = np.std([xgb_prob, lgb_prob, cat_prob]) * 2''')
    doc.add_paragraph(
        'Bootstrap — 100 выборок с возвратом, каждая обучена на 80% данных. '
        'Разброс моделей ×2 — эмпирический множитель для запаса.'
    )

    doc.add_heading('5.9. Price Targets', level=2)
    code(doc, '''def build_trade_plan(ticker, trend, ensemble, current_price, features):
    target = trend.get("target_price")
    entry = current_price
    take_profit = target if target else current_price * 1.05
    stop_loss = current_price * (1 - 0.02 - ensemble.get("uncertainty", 0.1) * 1.5)
    return {"current_price": current_price,
            "target_price": round(take_profit, 2),
            "stop_loss": round(stop_loss, 2),
            "potential_return_pct": round((take_profit/entry - 1) * 100, 1),
            "risk_pct": round((1 - stop_loss/entry) * 100, 1)}''')
    doc.add_paragraph(
        'stop_loss = entry × (1 - 0.02 - uncertainty×1.5). '
        'База 2% + запас на неопределённость. При высокой uncertainty '
        'стоп шире — модель не уверена, не надо выбивать позицию шумом.'
    )

    doc.add_page_break()

    # ═══════════════ 6. ВОЛАТИЛЬНОСТЬ ═══════════════
    doc.add_heading('6. Волатильность (VolatilityRegimeDetector)', level=1)
    doc.add_paragraph('src/analysis/volatility.py')

    doc.add_paragraph(
        'Определяет режим: LOW / NORMAL / HIGH / EXTREME. '
        'Влияет на веса в Fusion через adjustment — множители для источников сигнала.'
    )
    code(doc, '''def detect(self, df):
    if df.empty or len(df) < 20:
        return {"regime": "NORMAL", "adjustment": None}
    ret = df["close"].pct_change().dropna()
    hv = ret.rolling(20).std().iloc[-1] * (252 ** 0.5)
    atr = df["atr"].iloc[-1] if "atr" in df.columns else None
    if atr is None or len(df) < 60:
        return {"regime": "NORMAL", "adjustment": None, "hv": round(hv, 4)}
    atr_ratio = atr / df["atr"].iloc[-60:].mean()
    if hv > 0.6 or atr_ratio > 1.8:
        adj = {k: 0.7 for k in ("technical_mult","ml_mult","sentiment_mult")}
        return {"regime": "EXTREME", "adjustment": adj, ...}
    if hv > 0.4 or atr_ratio > 1.4:
        adj = {k: 0.85 for k in ("technical_mult","ml_mult","sentiment_mult")}
        return {"regime": "HIGH", "adjustment": adj, ...}
    if hv < 0.15 and atr_ratio < 0.7:
        adj = {"geo_mult": 1.2, "fundamental_mult": 1.1}
        return {"regime": "LOW", "adjustment": adj, ...}
    return {"regime": "NORMAL", "adjustment": None, ...}''')
    doc.add_paragraph(
        'Пороги: HV>60% — верхний квартиль MOEX (SBER: 25-45%). '
        'ATR_ratio>1.8 — два стандартных отклонения от средней ATR. '
        'Множители 0.7: при EXTREME тех. сигналы ненадёжны — ложные пробои. '
        'LOW: geo и фундамент получают больший вес (техника в консолидации бесполезна).'
    )

    doc.add_page_break()

    # ═══════════════ 7. MTF ═══════════════
    doc.add_heading('7. Мультивременной анализ (MTF)', level=1)
    doc.add_paragraph('src/analysis/multi_timeframe.py')

    doc.add_paragraph(
        'Согласованность сигналов на D1, W1, MN. Чем больше таймфреймов '
        'в одну сторону — тем выше mtf_signal. Вес в Fusion — 5% (акции) '
        'или 20% (облигации).'
    )
    code(doc, '''def analyze(self, df):
    if df.empty or len(df) < 200:
        return {"direction": 0.0, "agreement": 0.0, "details": {}}
    d1 = self._tf_signal(df)
    weekly = df.resample("W", on="date").agg(...)
    w1 = self._tf_signal(weekly) if len(weekly) > 50 else 0.0
    monthly = df.resample("ME", on="date").agg(...)
    mn = self._tf_signal(monthly) if len(monthly) > 12 else 0.0
    signals = {"D1": d1, "W1": w1, "MN": mn}
    non_zero = [s for s in signals.values() if s != 0.0]
    if not non_zero: return {"direction": 0.0, "agreement": 0.0, "details": signals}
    direction = np.mean(non_zero)
    agreement = 1.0 - np.std(non_zero)  # 1 = полное совпадение, 0 = хаос
    return {"direction": round(direction, 3), "agreement": round(max(0,agreement), 3)}''')
    doc.add_paragraph(
        'agreement = 1 - std(signals). '
        'Если D1=+1, W1=+0.8, MN=+0.6 → mean=0.8, std=0.16, agreement=0.84. '
        'Если D1=+1, W1=-0.5, MN=+0.2 → mean=0.23, std=0.61, agreement=0.39.'
    )

    doc.add_page_break()

    # ═══════════════ 8. СЕНТИМЕНТ ═══════════════
    doc.add_heading('8. Сентимент', level=1)
    doc.add_paragraph('src/collectors/news.py, src/collectors/sentiment.py, src/social/')

    doc.add_paragraph(
        'Оценка настроения [-1, 1] из новостей (RSS: TASS, RBC, Interfax) '
        'и LLM-анализа тональности (Groq → Ollama fallback). '
        'Вес в Fusion: 12% (акции) / 5% (облигации).'
    )

    doc.add_paragraph('Pipeline:')
    doc.add_paragraph(
        'NewsCollector → Groq (prompt «оцени тональность от -1 до 1») '
        '→ кластеризация (LDA) → аггрегация (соцсети 0.6, новости 0.4) '
        '→ divergence.'
    )
    code(doc, '''# LLM prompt (src/llm/prompts/analysis.py)
SENTIMENT_PROMPT = """
Оцени тональность финансовой новости от -1.0 до 1.0.
-1.0 = крайне негативная (дефолт, санкции, падение)
 0.0 = нейтральная (факты)
 1.0 = крайне позитивная (рост, дивиденды, контракты)
Верни ТОЛЬКО число.

Новость: {article_text}
"""

# Итоговый сигнал
sent_signal = raw_score * (1 - min(divergence, 0.5))''')
    doc.add_paragraph(
        'divergence = |social_score - news_score|. При сильном расхождении '
        'источников (0.5+) сигнал занижается вдвое. Cap 0.5 — даже при '
        'полном расхождении сигнал не обнуляется.'
    )

    doc.add_page_break()

    # ═══════════════ 9. ГЕОПОЛИТИКА ═══════════════
    doc.add_heading('9. Геополитический риск', level=1)
    doc.add_paragraph('src/geo/, src/data/geopolitical_risk_engine.py')

    doc.add_paragraph(
        'Шкала 0-10. Вес в Fusion: 17% (акции) — сопоставимо с фундаментом. '
        'Сигнал: geo_signal = -(geo_score / 10) → [-1, 0]. '
        'GEO_RISK_HIGH=7 → BUY → CAUTIOUS_BUY, позиция ≤10%.'
    )
    code(doc, '''class GeoRiskScorer:
    def score(self, news, events, macro) -> float:
        base = 3.0  # базовый риск РФ (не 0!)
        # sanctions_30d из EventFeatureBuilder
        if events.get("sanctions_30d", 0) > 0:
            base += events["sanctions_30d"] * 1.5
        # >30% негативных новостей
        if news.get("negative_sentiment_pct", 0) > 0.3:
            base += (news["negative_sentiment_pct"] - 0.3) * 5
        if macro.get("key_rate", 0) > 15: base += 1.0
        if macro.get("cpi", 0) > 8:       base += 1.0
        return min(base, 10.0)

    def adjust_for_sector(self, base, sector) -> float:
        mult = {"Нефть": 1.3, "Металлы": 1.2, "IT": 0.7, "Потреб": 0.8}
        return min(base * mult.get(sector, 1.0), 10.0)''')
    doc.add_paragraph(
        'Базовый риск 3.0 — структурный риск РФ (не обнуляем). '
        'Санкции: sanctions_30d — количество санкционных событий за 30 дней. '
        'Каждое даёт +1.5. Сектора: Нефть ×1.3 (экспортные санкции), IT ×0.7.'
    )

    doc.add_page_break()

    # ═══════════════ 10. FUSION ENGINE ═══════════════
    doc.add_heading('10. Fusion Engine', level=1)
    doc.add_paragraph('src/signals/engine.py | SignalFusionEngine.fuse()')

    doc.add_paragraph(
        'Центральный узел: взвешивает 6 источников сигнала, добавляет '
        'макро-коррекцию, тренд, штраф за события -> BUY/SELL/HOLD.'
    )

    doc.add_heading('10.1. Веса', level=2)
    code(doc, '''BASE_WEIGHTS = {
    "technical":   0.35,   # MOEX — тех. драйвер
    "fundamental": 0.18,   # якорь
    "geo":         0.17,   # РФ чувствительность
    "ml":          0.13,   # усилитель
    "sentiment":   0.12,   # фильтр новостей
    "mtf":         0.05,   # подтверждение
}
BOND_WEIGHTS = {
    "technical":   0.10, "fundamental": 0.40, "geo": 0.15,
    "ml":          0.10, "sentiment":   0.05, "mtf": 0.20,
}''')
    doc.add_paragraph(
        'Почему tech=0.35: MOEX — технический рынок (низкая ликвидность, '
        'сильная роль уровней). Баланс подобран экспертно, не оптимизирован.'
    )

    doc.add_heading('10.2. Risk-профили', level=2)
    code(doc, '''RISK_PROFILES = {
    "conservative": {"weights": {"tech":0.30,"fund":0.25,"geo":0.20,
                     "ml":0.08,"sent":0.07,"mtf":0.10},
                     "max_pos":10, "min_conf":0.4, "geo_thr":6.0},
    "balanced":     {"weights": {"tech":0.35,"fund":0.18,"geo":0.17,
                     "ml":0.13,"sent":0.12,"mtf":0.05},
                     "max_pos":20, "min_conf":0.3, "geo_thr":7.0},
    "aggressive":   {"weights": {"tech":0.40,"fund":0.10,"geo":0.10,
                     "ml":0.20,"sent":0.15,"mtf":0.05},
                     "max_pos":35, "min_conf":0.2, "geo_thr":8.0},
}''')
    doc.add_paragraph(
        'Горизонт корректирует веса: long → fund×1.3, tech×0.8. '
        'Short → tech×1.3, fund×0.7. После коррекции — нормировка на 1.0.'
    )

    doc.add_heading('10.3. Динамические коррекции', level=2)
    doc.add_paragraph('Негативный тех. сигнал → boost tech +0.10 за счёт sent/mtf/geo:')
    code(doc, '''if tech_score_raw < 0:
    weights["technical"] += 0.10
    for k in ["sentiment", "mtf", "geo"]:
        weights[k] -= 0.04
    # нормировка'''[:1])
    doc.add_paragraph('Волатильность (из раздела 6) множит веса:')
    code(doc, '''if vol_regime and vol_regime.get("adjustment"):
    for key in weights:
        mk = f"{key}_mult"
        if mk in vol_regime["adjustment"]:
            weights[key] *= vol_regime["adjustment"][mk]
    total = sum(weights.values())
    for k in weights: weights[k] /= total''')
    doc.add_paragraph('Тренд и события:')
    code(doc, '''trend_adjustment = trend_slope * 0.08 * trend_strength
if trend_changed and trend_slope < 0:
    trend_adjustment -= 0.05
event_penalty = event_risk_score * 0.05
if sanctions_spike:  event_penalty += 0.03''')
    doc.add_paragraph(
        'Множитель 0.08 — эмпирика: макс. влияние тренда ~0.08 при slope=1. '
        'Санкционный штраф +0.03 суммируется с базовым 0.05 за событие.'
    )

    doc.add_heading('10.4. Макро', level=2)
    code(doc, '''MACRO_THRESHOLDS = {
    "brent":     {"high": 80,  "high_adj": 0.03,  "low": 50, "low_adj": -0.05},
    "key_rate":  {"high": 15,  "high_adj": -0.05, "low": 7,  "low_adj": 0.03},
    "cpi":       {"high": 8,   "high_adj": -0.04, "low": 4,  "low_adj": 0.02},
    "ofz_10y":   {"high": 12,  "high_adj": -0.03, "low": 6,  "low_adj": 0.02},
    "m2":        {"high": 70_000, "high_adj": 0.02, "low": 50_000, "low_adj": -0.02},
    "imoex":     {"high": 3500, "high_adj": 0.02, "low": 2500, "low_adj": -0.03},
}
MACRO_MAX_ADJUSTMENT = 0.10''')
    doc.add_paragraph(
        'Brent >80 = профицит бюджета (+0.03). Key rate >15 = жёсткая ДКП (-0.05). '
        'Лимит 0.10: макро не переопределяет сигналы. '
        'Пороги — исторические уровни (2014-2026).'
    )

    doc.add_heading('10.5. Bond-специфичный scoring', level=2)
    code(doc, '''if is_bond and bond_offering:
    bond_signal = analyze_bond(bond_offering, key_rate=key_rate)
    bond_score = bond_signal.get("score", 0.0)
    weighted_score = bond_score * weights["fundamental"] * 2 + weighted_score * 0.5
    fundamental["risk"] = max(fundamental.get("risk", 0.5), bond_risk)''')
    doc.add_paragraph(
        'bond_score × 2 — усиливаем вес облигационного сигнала, '
        'т.к. фундаментальный анализ для бондов точнее тех. '
        'Исходный weighted_score тоже учитывается (×0.5).'
    )

    doc.add_heading('10.6. Итоговый сигнал', level=2)
    code(doc, '''# Сигналы
tech_signal  = technical.get("score", 0.0)     # из TechnicalAnalyzer
fund_signal  = (1 - fund_risk) * 2 - 1          # [-1, 1]
geo_signal   = -(geo_score / 10)                # [-1, 0]
ml_signal    = ml_prediction.get("signal_score", 0.0)
sent_signal  = sentiment["score"] * (1 - min(divergence, 0.5))
mtf_signal   = mtf["direction"]

weighted_score = (
    tech_score  * weights["technical"]
    + fund_signal * weights["fundamental"]
    + geo_signal  * weights["geo"]
    + ml_signal   * weights["ml"]
    + sent_signal * weights["sentiment"]
    + mtf_signal  * weights["mtf"]
    + macro_adjustment * MACRO_MAX_ADJUSTMENT
    + trend_adjustment * weights["ml"]
    - event_penalty
)

if weighted_score > 0.02:    action = "BUY"
elif weighted_score < -0.02:  action = "SELL"
else:                         action = "HOLD"

# Confidence с risk-adj
confidence = abs(weighted_score) / max_absolute
risk_adj = 1.0 + sharpe*0.05 + sortino*0.03 - mdd*2 + calmar*0.02 + omega*0.01
confidence *= max(risk_adj, 0.3)

# Фильтры ложных сигналов
if action == "BUY" and bearish_smas >= 2 and tech_score < 0:
    action = "HOLD"   # тех. фильтр
if action == "BUY" and trend_adjustment < -0.02 and ml_signal < 0.1:
    action = "HOLD"   # нисходящий тренд

max_pct = BASE_POSITION_PCT.get(action, 10)
if geo_risk > GEO_RISK_HIGH:    max_pct = min(max_pct, 10)
if fund_risk > FUND_RISK_HIGH:  max_pct = min(max_pct, 10)''')
    doc.add_paragraph(
        "Пороги 0.02/-0.02 — эмпирические, не валидированы walk-forward\'ом "
        '(в отличие от тех. порога 0.20). Это потенциальное улучшение.'
    )

    doc.add_page_break()

    # ═══════════════ 11-21. ОСТАЛЬНЫЕ МОДУЛИ ═══════════════
    # Сжато: алерты, торговля, портфель, уведомления, LLM, интерфейсы,
    # БД, соцсети, аномалии, планировщик, фронтенд

    doc.add_heading('11. Алерты', level=1)
    doc.add_paragraph(
        'src/alerts/. AlertEngine: новости → дедупликация (по заголовку + '
        'временное окно) → детекция аномалий (AnomalyDetector) → '
        'приоритизация (важность 0-1) → push (email/TG/webpush).'
    )
    add_table(doc,
        ['Файл', 'Класс', 'Назначение'],
        [
            ('engine.py', 'AlertEngine', 'оркестратор'),
            ('deduplicator.py', 'AlertDeduplicator', 'дедупликация'),
            ('scorer.py', 'build_alert', 'важность алерта'),
            ('prioritizer.py', 'AlertPrioritizer', 'сортировка'),
            ('generators.py', 'AlertGenerators', 'типы событий'),
            ('push.py', 'PushDispatcher', 'отправка'),
            ('smart.py', 'SmartAlertEngine', 'адаптивные алерты (ML)'),
        ]
    )

    doc.add_page_break()

    doc.add_heading('12. Торговля', level=1)
    doc.add_paragraph('src/trading/')

    doc.add_paragraph(
        '5 брокеров: Т-Банк (Tinkoff), Alor, БКС, Финам, OpenAPI. '
        'BaseBroker: get_portfolio, place_order, cancel_order, get_positions, get_orderbook.'
    )
    code(doc, '''class ExecutionEngine:
    async def execute(self, order, mode):
        check_order_aml(order)          # AML
        check_position_limit(order)     # лимиты
        check_short_eligibility(order)  # шорт-лист
        if not get_circuit_breaker(ticker).ready():
            raise CircuitBreakerOpenError
        broker = self._get_broker(order.broker)
        result = await broker.place_order(order)
        save_order(result)
        await _notify_trade(result)
        return result''')
    doc.add_paragraph(
        'Режимы: DRY_RUN (симуляция), MANUAL (подтверждение), AUTO (авто). '
        'Margin status: safe → warning → margin_call → liquidation. '
        'Short rates: SBER 15%, GAZP 20%, LKOH 12%, VTBR 25%. '
        'AML: round_trip, structuring, velocity_anomaly.'
    )

    doc.add_page_break()

    doc.add_heading('13. Портфель', level=1)
    doc.add_paragraph('src/portfolio/')

    doc.add_paragraph(
        'AllocationEngine: тиры капитала (≤1k → 1 поз, ≤3k → 2 поз, >3k → секторные лимиты). '
        'Секторные лимиты: Нефть 35%, Банки 25%, Финансы 20%. '
        'Макс. 15 бумаг, лимит pick\'ов зависит от капитала.'
    )
    code(doc, '''def compute_risk_metrics(price_series):
    returns = np.diff(arr) / arr[:-1]
    return {"sharpe": _sharpe_ratio(returns), "sortino": ...,
            "max_drawdown": ..., "calmar": ..., "omega": ...}
# Sharpe  = mean(excess) / std(returns) * sqrt(252)
# Sortino = mean(excess) / std(downside) * sqrt(252)
# Calmar  = CAGR / max_drawdown
# Omega   = sum(gains) / abs(sum(losses))''')

    doc.add_page_break()

    doc.add_heading('14. Уведомления', level=1)
    doc.add_paragraph(
        'src/notifications/. Каналы: email (SMTP), Telegram (bot), WebPush (VAPID). '
        'Jinja2-шаблоны: email/(alert, daily, signal), telegram/, webpush/. '
        'Retry: 3 попытки с exponential backoff (2^attempt сек).'
    )

    doc.add_page_break()

    doc.add_heading('15. LLM', level=1)
    doc.add_paragraph(
        'src/llm/. Groq (основной) → Ollama (fallback). Rate limiter (токены/мин). '
        'Промпты: анализ (src/llm/prompts/analysis.py), вопрос, отчёт. '
        'Инструменты: WolframAlpha (src/llm/tools/wolfram.py).'
    )

    doc.add_page_break()

    doc.add_heading('16. CLI и API', level=1)
    doc.add_paragraph(
        'Typer CLI (main.py, src/cli/): finn init, update TICKER, analyze TICKER, '
        'list-instruments, rates, portfolio, news TICKER.'
    )
    doc.add_paragraph(
        'FastAPI (src/interfaces/api/server.py, порт 8000): '
        '/api/v1/auth/*, /instruments, /analysis/{ticker}, /backtest/*, '
        '/portfolio/*, /bonds/*, /alerts/preferences, /trading/*, /health.'
    )
    doc.add_paragraph(
        'Telegram bot (python-telegram-bot): /start, /analyze, /portfolio, /news, /bonds. '
        'NLQ (src/interfaces/nlq/): свободные запросы → структурированные.'
    )

    doc.add_page_break()

    doc.add_heading('17. База данных', level=1)
    doc.add_paragraph(
        'PostgreSQL + SQLAlchemy async (asyncpg). 29 Alembic-миграций. '
        'PgBouncer (deploy/pgbouncer.ini). Grafana-дашборд (deploy/grafana/). '
        'SQLite (data/finn.db) для dev/тестов.'
    )
    add_table(doc,
        ['Модель', 'Назначение'],
        [
            ('Instrument', 'акции/облигации'),
            ('Price', 'OHLCV'),
            ('Indicator', 'тех. индикаторы'),
            ('Dividend', 'дивиденды'),
            ('FundamentalMetric', 'P/E, P/B, ROE, D/E'),
            ('BondOffering', 'облигации'),
            ('News', 'новости'),
            ('Signal', 'Fused сигналы'),
            ('Portfolio', 'портфели'),
            ('User', 'пользователи'),
            ('BrokerCredential', 'брокеры'),
            ('AuditLog', 'аудит'),
        ]
    )

    doc.add_page_break()

    doc.add_heading('18. Социальный сентимент', level=1)
    doc.add_paragraph(
        'src/social/. Анализ Telegram-каналов + RSS. '
        'Аггрегация: соцсети 0.6, новости 0.4. '
        'SentimentAggregator: среднее по источникам, divergence = |social - news|.'
    )

    doc.add_page_break()

    doc.add_heading('19. Детекция аномалий', level=1)
    doc.add_paragraph('src/analysis/anomaly/')
    add_table(doc,
        ['Файл', 'Назначение'],
        [
            ('detector.py', 'главный детектор (train_all + detect)'),
            ('autoencoder.py', 'нейросетевой детектор (PyTorch)'),
            ('volume_anomaly.py', 'аномалии объёмов'),
            ('sentiment_anomaly.py', 'аномалии сентимента'),
            ('topic_anomaly.py', 'аномалии тем'),
            ('source_anomaly.py', 'аномалии источников'),
        ]
    )

    doc.add_page_break()

    doc.add_heading('20. Планировщик', level=1)
    doc.add_paragraph('Celery worker (src/tasks/worker.py):')
    add_table(doc,
        ['Задача', 'Расписание'],
        [
            ('collect_data', '15 мин'),
            ('update_ml_models', 'ежедневно'),
            ('generate_daily_report', '19:00 ежедневно'),
            ('check_alerts', '5 мин'),
            ('retrain_ensemble', 'еженедельно'),
        ]
    )
    doc.add_paragraph('Windows: run_daily.ps1 (обёртка для uv run finn update).')

    doc.add_page_break()

    doc.add_heading('21. Web-фронтенд', level=1)
    doc.add_paragraph(
        'Next.js 16, TypeScript, App Router, TanStack Query, Recharts, '
        'Tailwind CSS + shadcn/ui.'
    )
    add_table(doc,
        ['Директория', 'Назначение'],
        [
            ('app/ (auth)', 'login, register'),
            ('app/ (protected)', 'alerts, dashboard, instruments, paper, portfolio'),
            ('components/', 'AlertList, CandlestickChart, PortfolioOverview и др.'),
            ('features/', 'bonds, market, portfolio (feature-based)'),
            ('hooks/', 'useAuth, useBond, usePortfolio'),
            ('types/', 'TypeScript (bond, portfolio, chart)'),
        ]
    )

    doc.add_page_break()

    # ═══════════════ 22. КОНТРАКТЫ ═══════════════
    doc.add_heading('22. Контракты данных', level=1)
    doc.add_paragraph('Форматы входа/выхода основных анализаторов.')

    doc.add_heading('22.1. TechnicalAnalyzer.compute_all() вход', level=2)
    add_table(doc,
        ['Поле', 'Тип', 'Описание'],
        [
            ('date', 'datetime64', 'дата'),
            ('open', 'float64', 'открытие, руб'),
            ('high', 'float64', 'максимум, руб'),
            ('low', 'float64', 'минимум, руб'),
            ('close', 'float64', 'закрытие, руб'),
            ('volume', 'int64', 'объём, шт'),
        ]
    )
    doc.add_paragraph('Источник: MOEX ISS. Периодичность: ежедневно.')

    doc.add_heading('22.2. FundamentalAnalyzer.analyze() — metrics', level=2)
    add_table(doc,
        ['Поле', 'Тип', 'Описание', 'Источник'],
        [
            ('market_cap', 'float', 'капитализация, руб', 'MOEX'),
            ('pe_ratio', 'float', 'P/E (<0 = убыток)', 'MOEX'),
            ('pb_ratio', 'float', 'P/B', 'MOEX'),
            ('roe', 'float', 'ROE, % (напр. 15.0)', 'MOEX'),
            ('eps', 'float', 'прибыль на акцию, руб', 'MOEX'),
            ('debt_equity', 'float', 'D/E', 'MOEX'),
        ]
    )
    doc.add_paragraph(
        'Если поле отсутствует — проверка пропускается (не штраф). '
        'Периодичность: ежеквартально.'
    )

    doc.add_heading('22.3. FundamentalAnalyzer.analyze_bond() — bond_offering', level=2)
    add_table(doc,
        ['Поле', 'Тип', 'Описание'],
        [
            ('yield_to_maturity', 'float', 'YTM, % (напр. 8.5)'),
            ('credit_rating', 'str', 'рейтинг: "BB+" или "ruAAA"'),
            ('duration_years', 'float', 'дюрация, лет'),
        ]
    )
    doc.add_paragraph(
        'Рейтинг: международный (S&P: AAA→D) или российский (АКРА: ruAAA→ruD). '
        'Префиксный матчинг.'
    )

    doc.add_heading('22.4. SignalFusionEngine.fuse() выход', level=2)
    add_table(doc,
        ['Поле', 'Тип', 'Диапазон'],
        [
            ('action', 'str', 'BUY/SELL/HOLD/CAUTIOUS_BUY/NEUTRAL'),
            ('confidence', 'float', '[0, 1]'),
            ('weighted_score', 'float', '[-1, 1]'),
            ('reasons', 'list[str]', '≤8 шт'),
            ('max_portfolio_pct', 'int', '[5, 50]'),
            ('components.technical.score', 'float', '[-1, 1]'),
            ('components.fundamental_risk', 'float', '[0, 1]'),
            ('components.geo_risk', 'float', '[0, 10]'),
            ('components.ml.signal_score', 'float', '[-1, 1]'),
            ('components.sentiment.score', 'float', '[-1, 1]'),
            ('components.mtf.direction', 'float', '[-2, 2]'),
        ]
    )

    doc.add_page_break()

    # ═══════════════ 23. ДЕГРАДАЦИЯ ═══════════════
    doc.add_heading('23. Деградация', level=1)
    doc.add_paragraph('Поведение при неполных данных.')

    doc.add_heading('23.1. Тех. анализ', level=2)
    doc.add_paragraph(
        'df пустой → risk=0.5, signal="недостаточно данных". '
        '<50 дней → action=NEUTRAL, confidence=0. '
        'SMA-колонки с NaN — пропускаются без штрафа.'
    )

    doc.add_heading('23.2. Фундаментальный', level=2)
    doc.add_paragraph(
        'prices пустой → risk=0.5. metrics=None → все мультипликаторы '
        'пропускаются (штатно). dividends пустой → без дивидендного блока.'
    )

    doc.add_heading('23.3. ML', level=2)
    doc.add_paragraph(
        'Данных < min_train_rows → None (ml_signal=0, confidence=0). '
        'Ансамбль не обучен → fallback на XGBoost, если есть.'
    )

    doc.add_heading('23.4. Сентимент', level=2)
    doc.add_paragraph(
        'Новостей нет → sentiment={"score":0, "source":"none", "divergence":0}. '
        'Groq недоступен → Ollama. Оба недоступны → sent_signal=0.'
    )

    doc.add_heading('23.5. Гео', level=2)
    doc.add_paragraph(
        'Нет событий → sanctions_30d=0. Нет макро → проверки пропускаются. '
        'Минимум geo_score=3.0 (консервативно).'
    )

    doc.add_heading('23.6. Fusion', level=2)
    doc.add_paragraph(
        'После коррекций все веса нормируются на 1.0. '
        'Если все сигналы=0 → weighted_score = macro + trend. '
        'Если и макро нет → HOLD, confidence=0.'
    )

    doc.add_page_break()

    # ═══════════════ 24. ПРОБЛЕМЫ ═══════════════
    doc.add_heading('24. Известные проблемы', level=1)

    doc.add_heading('24.1. tech_signal vs fund_signal', level=2)
    doc.add_paragraph(
        'В старой версии документации встречалось: '
        'tech_signal = (1 - fund_risk) * 2 - 1 (обе формулы ссылаются на fund_risk). '
        'Это баг копипаста. В коде (src/signals/engine.py): '
        'tech_signal = technical.get("score", 0.0), fund_signal = (1 - fund_risk) * 2 - 1.'
    )

    doc.add_heading('24.2. MACD whipsaw', level=2)
    doc.add_paragraph(
        'Порог max(0.01, macd_hist.std() * 0.1) эмпирический. '
        'На VTBR std может быть слишком высок — сигналы пропускаются. '
        'Лучше: адаптивный перцентиль.'
    )

    doc.add_heading('24.3. Sector benchmarks', level=2)
    doc.add_paragraph(
        'Статические медианы (данные 2024-2025). IT-сектор быстро меняется — '
        'нужен пересмотр. План: авто-расчёт из БД.'
    )

    doc.add_heading('24.4. Fusion пороги', level=2)
    doc.add_paragraph(
        '0.02/-0.02 для action — эмпирика, не WF-валидирована. '
        'В отличие от тех. порога 0.20.'
    )

    doc.add_heading('24.5. Meta-learner', level=2)
    doc.add_paragraph(
        'Discard при accuracy<0.52. На коротких историях meta часто '
        'отбрасывается. Влияние на ensemble без meta не измерено.'
    )

    doc.add_heading('24.6. Тестовое покрытие', level=2)
    doc.add_paragraph(
        'Technical: test_technical.py (основное). '
        'Fundamental: test_fundamental.py (12 тестов). '
        'Fusion: test_signal.py. ML: test_ml.py, test_price_targets.py. '
        'Алерты: 6 файлов. '
        'Не покрыто: точное совпадение весов после коррекций, '
        'комбинация факторов деградации, MTF на пустых данных.'
    )

    doc.add_page_break()

    # ═══════════════ 25. ПРИМЕР ═══════════════
    doc.add_heading('25. Пример сквозной', level=1)
    doc.add_paragraph('Данные: SBER за 2024-2026 (последние 5 дней).')

    doc.add_paragraph('Исходные OHLCV (последние 5 строк):')
    code(doc, '''date        open    high    low     close   volume
2026-07-10  312.0   315.0   310.5   313.2   12_450_000
2026-07-13  313.5   316.2   311.0   314.8   11_800_000
2026-07-14  314.5   317.0   313.0   316.1   14_200_000
2026-07-15  316.0   318.5   314.5   315.3   10_900_000
2026-07-16  315.0   316.8   312.0   313.8   13_500_000''')

    doc.add_paragraph('Индикаторы (последняя строка):')
    code(doc, '''RSI(14)=52.3   → нейтрально
MACD hist=+0.15 → слабый бычий
SMA20=311.5    → close(313.8) > SMA20, SMA50(305.2), SMA200(290.1) → +1.5
BB touch       → close внутри полос → 0
Volume ratio=1.12 → норма
1d return=-0.48% → 0

score=2.5, max_score=4.5, normalized=0.56 → BUY''')

    doc.add_paragraph('Фундамент:')
    code(doc, '''sector=Банки, P/E=7.2 (сектор 6.0), ROE=18.5%, D/E=3.2 (сектор 5.0)
MCAP=3.5 трлн → крупная
risk=0.0, fund_signal=+1.0''')

    doc.add_paragraph('Fusion:')
    code(doc, '''Веса (balanced): 0.35/0.18/0.17/0.13/0.12/0.05
tech=+0.56 fund=+1.0 geo=-0.3 ml=+0.20 sent=+0.1 mtf=+0.5
weighted=0.56*0.35 + 1.0*0.18 + (-0.3)*0.17 + 0.20*0.13 + 0.1*0.12 + 0.5*0.05
       = 0.196 + 0.18 - 0.051 + 0.026 + 0.012 + 0.025 = 0.388
macro: Brent=85 (high +0.03), key_rate=16 (high -0.05) → -0.02
trend=+0.01, events=0
Итог: 0.388 - 0.002 + 0.001 = 0.387 → BUY
confidence = 0.387 / 0.75 = 0.52
max_pos = 20% (balanced)''')

    doc.add_page_break()

    # ═══════════════ 26. ТАБЛИЦЫ ═══════════════
    doc.add_heading('26. Таблицы порогов', level=1)

    doc.add_heading('26.1. Технический анализ', level=2)
    add_table(doc,
        ['Параметр', 'Значение', 'Файл', 'Обоснование'],
        [
            ('RSI период', '14', 'technical.py', 'стандарт Wilder'),
            ('RSI oversold', '<30', 'technical.py', 'стандарт'),
            ('RSI overbought', '>70', 'technical.py', 'стандарт'),
            ('MACD', '12/26/9', 'technical.py', 'стандарт Appel'),
            ('BB', '20/2σ', 'technical.py', 'стандарт Bollinger'),
            ('SMA', '20/50/200', 'technical.py', 'месяц/квартал/год'),
            ('Score BUY', '>0.20', 'technical.py', 'WF на 10 тикеров: OOS 0.553'),
            ('Whipsaw', 'max(0.01, std*0.1)', 'technical.py', 'эмпирика'),
        ], col_widths=[3, 2.5, 2, 9]
    )

    doc.add_paragraph('')
    doc.add_heading('26.2. Фундаментальный', level=2)
    add_table(doc,
        ['Параметр', 'Значение', 'Файл', 'Обоснование'],
        [
            ('MCAP low', '1 млрд руб', 'fundamental.py', 'малая кап на MOEX'),
            ('MCAP high', '100 млрд руб', 'fundamental.py', 'blue chip'),
            ('P/E mult', '3× сектор', 'fundamental.py', 'экспертно'),
            ('ROE mult', '0.3× сектор', 'fundamental.py', 'экспертно'),
            ('D/E mult', '2× сектор', 'fundamental.py', 'экспертно'),
            ('Годовое падение', '<-30%', 'fundamental.py', '2σ MOEX'),
            ('Волатильность', '>50% годовых', 'fundamental.py', 'верхний квартиль MOEX'),
            ('Bond спред', '>5%', 'fundamental.py', '500 б.п. → дефолтный риск'),
            ('Bond duration', '>10 лет', 'fundamental.py', 'чувствительность 10%/1% ставки'),
            ('Рейтинг', '0.00-0.30', 'fundamental.py', 'S&P + АКРА'),
        ], col_widths=[3, 2.5, 2, 9]
    )

    doc.add_paragraph('')
    doc.add_heading('26.3. ML', level=2)
    add_table(doc,
        ['Параметр', 'Значение', 'Файл', 'Обоснование'],
        [
            ('lookahead', '5 дней', '_base.py', 'краткосрок'),
            ('threshold', '3% (адаптивный)', 'walk_forward.py', '≈2× спред MOEX'),
            ('признаков', '22+', '_base.py', '14 баз + 4 события + 7 макро'),
            ('OOS min', '0.52', 'walk_forward.py', 'выше random'),
            ('WF folds', '3', 'walk_forward.py', 'Pesaran-Timmermann'),
            ('WF gap', '20 дней', 'walk_forward.py', '≈месяц, защита автокорр'),
            ('Meta C', '0.5', 'ensemble.py', 'L2 рег.'),
            ('Вес модели', '(acc-0.5)*4', 'walk_forward.py', '0.5→0, 0.75→1'),
        ], col_widths=[3, 2.5, 2, 9]
    )

    doc.add_paragraph('')
    doc.add_heading('26.4. Fusion', level=2)
    add_table(doc,
        ['Параметр', 'Значение', 'Файл', 'Обоснование'],
        [
            ('BASE_WEIGHTS', '0.35/0.18/0.17/0.13/0.12/0.05', 'engine.py', 'экспертно для MOEX'),
            ('BOND_WEIGHTS', '0.10/0.40/0.15/0.10/0.05/0.20', 'engine.py', 'акцент на кредит'),
            ('BUY порог', '>0.02', 'engine.py', 'эмпирика'),
            ('GEO_RISK_HIGH', '7.0', 'constants.py', 'верхний дециль'),
            ('FUND_RISK_HIGH', '0.6', 'constants.py', '3+ аномалии'),
            ('MACRO_MAX', '0.10', 'constants.py', 'макро не переопределяет'),
        ], col_widths=[3, 2.5, 2, 9]
    )

    doc.add_paragraph('')
    doc.add_heading('26.5. Макро', level=2)
    add_table(doc,
        ['Индикатор', 'High', 'Adj', 'Low', 'Adj', 'Обоснование'],
        [
            ('Brent', '80', '+0.03', '50', '-0.05', 'бюджет РФ ~60'),
            ('Ключевая', '15%', '-0.05', '7%', '+0.03', 'исторический диапазон ЦБ'),
            ('CPI', '8%', '-0.04', '4%', '+0.02', 'цель ЦБ=4%'),
            ('ОФЗ 10Y', '12%', '-0.03', '6%', '+0.02', 'исторический'),
            ('IMOEX', '3500', '+0.02', '2500', '-0.03', 'исторический'),
        ], col_widths=[2, 1.5, 1, 1.5, 1, 9.5]
    )

    doc.add_paragraph('')
    doc.add_heading('26.6. Торговля', level=2)
    add_table(doc,
        ['Параметр', 'Значение', 'Файл', 'Обоснование'],
        [
            ('Short SBER', '15%', 'constants.py', 'ставка займа'),
            ('Short GAZP', '20%', 'constants.py', ''),
            ('Налог дивиденды', '13%', 'constants.py', 'НДФЛ'),
            ('Налог долгосроч. (3+ лет)', '0%', 'constants.py', 'льгота РФ'),
            ('Комиссия брокера', '0.04%', 'constants.py', 'Т-Банк средняя'),
        ], col_widths=[3, 2.5, 2, 9]
    )

    # ── Footer ──
    doc.add_paragraph('')
    try:
        git_hash = subprocess.check_output(
            ['git', 'rev-parse', '--short', 'HEAD'],
            cwd='D:\\finn-help', text=True
        ).strip()[:7]
    except Exception:
        git_hash = 'unknown'

    doc.add_paragraph(
        '---\n'
        f'Commit: {git_hash} | {datetime.date.today().strftime("%d.%m.%Y")}\n'
        'Генерация: python docs/generate_docs.py | python-docx\n'
        'Правка кода → перегенерируй этой же командой.'
    )

    out_path = 'D:\\finn-help\\docs\\FinAdvisor_Technical_Documentation.docx'
    doc.save(out_path)
    print(f'Saved: {out_path}')


if __name__ == '__main__':
    build()
